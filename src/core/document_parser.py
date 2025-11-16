"""
Document Parser - Extract text content from various document formats
"""

import os
import PyPDF2
import docx
import pandas as pd
from typing import Dict, Any, List

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    print("pdfplumber not available - using PyPDF2 only")

try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image
    OCR_AVAILABLE = True
    
    # Try to set Tesseract path (Windows)
    import shutil
    tesseract_path = shutil.which('tesseract')
    if tesseract_path:
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
        print(f"Found Tesseract in PATH: {tesseract_path}")
    else:
        # Check local directory first, then common installation paths
        current_dir = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
        possible_paths = [
            os.path.join(current_dir, 'tesseract.exe'),  # Local directory
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            r'C:\Users\{}\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'.format(os.environ.get('USERNAME', ''))
        ]
        for path in possible_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                print(f"Found Tesseract at: {path}")
                break
        else:
            print("Tesseract not found in common locations - OCR will be limited")
            
except ImportError:
    OCR_AVAILABLE = False
    print("OCR libraries not available - install pytesseract and pdf2image for OCR support")

class DocumentParser:
    """Parse various document formats and extract text content."""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_doc,
            '.xlsx': self._parse_xlsx,
            '.xls': self._parse_xls,
            '.txt': self._parse_txt
        }
    
    def parse_file(self, filepath: str) -> str:
        """
        Parse a file and return its text content.
        
        Args:
            filepath: Path to the file to parse
            
        Returns:
            Extracted text content as string
            
        Raises:
            ValueError: If file format is not supported
            IOError: If file cannot be read
        """
        if not os.path.exists(filepath):
            raise IOError(f"File not found: {filepath}")
            
        _, ext = os.path.splitext(filepath.lower())
        
        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {ext}")
            
        try:
            return self.supported_formats[ext](filepath)
        except Exception as e:
            raise IOError(f"Failed to parse {filepath}: {str(e)}")
    
    def _parse_pdf(self, filepath: str) -> str:
        """Parse PDF file and extract text using multiple methods."""
        text_content = []
        extraction_method = "Unknown"
        
        # Method 1: Try PyPDF2 first
        try:
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                print(f"PDF has {len(pdf_reader.pages)} pages")
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    print(f"PyPDF2 - Page {page_num + 1} extracted {len(page_text)} characters")
                    
                    if page_text.strip():
                        text_content.append(page_text)
                
                if text_content:
                    extraction_method = "PyPDF2"
                    full_text = "\n".join(text_content)
                    print(f"PyPDF2 successfully extracted {len(full_text)} characters total")
                    return full_text
                    
        except Exception as e:
            print(f"PyPDF2 failed: {str(e)}")
        
        # Method 2: Try pdfplumber if PyPDF2 failed or produced no content
        if PDFPLUMBER_AVAILABLE:
            try:
                print("Trying pdfplumber extraction...")
                text_content = []
                
                with pdfplumber.open(filepath) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                print(f"pdfplumber - Page {page_num + 1} extracted {len(page_text)} characters")
                                text_content.append(page_text)
                            else:
                                print(f"pdfplumber - Page {page_num + 1} extracted 0 characters")
                        except Exception as page_error:
                            print(f"pdfplumber - Page {page_num + 1} failed: {str(page_error)}")
                            continue
                
                if text_content:
                    extraction_method = "pdfplumber"
                    full_text = "\n".join(text_content)
                    print(f"pdfplumber successfully extracted {len(full_text)} characters total")
                    return full_text
                    
            except Exception as e:
                print(f"pdfplumber failed: {str(e)}")
        
        # Method 3: Try OCR if both text extraction methods failed
        if OCR_AVAILABLE:
            try:
                print("Trying OCR extraction (this may take a moment)...")
                return self._parse_pdf_with_ocr(filepath)
            except Exception as e:
                print(f"OCR extraction failed: {str(e)}")
        
        # If all methods failed or produced no content
        try:
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
        except:
            num_pages = "Unknown"
        
        ocr_status = "OCR attempted but failed" if OCR_AVAILABLE else "OCR not available (Tesseract not installed)"
        
        return f"""PDF text extraction produced minimal content.

File: {os.path.basename(filepath)}
Pages: {num_pages}
Methods tried: PyPDF2{', pdfplumber' if PDFPLUMBER_AVAILABLE else ''}{', OCR' if OCR_AVAILABLE else ''}
Status: {ocr_status}

This PDF likely contains:
- Scanned images rather than searchable text
- Complex formatting or embedded graphics
- Password protection or encryption

Next steps:
1. Install Tesseract OCR for automatic text recognition from images
2. Try converting the PDF to text using Adobe Reader or online tools
3. Check if the PDF has selectable text (try copying text manually)
4. Convert to Word or other text-based format

OCR Setup (for scanned documents):
- Download Tesseract OCR from: https://github.com/UB-Mannheim/tesseract/wiki
- Install to default location (C:\\Program Files\\Tesseract-OCR\\)
- Restart the application to enable OCR functionality

If this PDF should contain searchable text, please check:
- The PDF is not corrupted
- You have the latest version of the document
- The PDF creator used text-based formatting (not image scans)
"""
    
    def _parse_docx(self, filepath: str) -> str:
        """Parse DOCX file and extract text."""
        try:
            doc = docx.Document(filepath)
            text_content = []
            
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
                        
        except Exception as e:
            raise IOError(f"Error parsing DOCX: {str(e)}")
            
        return "\n".join(text_content)
    
    def _parse_doc(self, filepath: str) -> str:
        """Parse DOC file (legacy Word format)."""
        # For now, return a placeholder - would need python-docx2txt or similar
        return f"DOC file parsing not fully implemented: {os.path.basename(filepath)}"
    
    def _parse_xlsx(self, filepath: str) -> str:
        """Parse Excel XLSX file and extract text."""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(filepath)
            text_content = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(filepath, sheet_name=sheet_name)
                
                # Add sheet name as header
                text_content.append(f"=== Sheet: {sheet_name} ===")
                
                # Convert DataFrame to text
                # Replace NaN with empty string
                df_cleaned = df.fillna('')
                
                # Add column headers
                headers = " | ".join(str(col) for col in df_cleaned.columns)
                text_content.append(headers)
                text_content.append("-" * len(headers))
                
                # Add rows
                for _, row in df_cleaned.iterrows():
                    row_text = " | ".join(str(value) for value in row.values)
                    text_content.append(row_text)
                
                text_content.append("")  # Add blank line between sheets
                
        except Exception as e:
            raise IOError(f"Error parsing Excel file: {str(e)}")
            
        return "\n".join(text_content)
    
    def _parse_xls(self, filepath: str) -> str:
        """Parse Excel XLS file (legacy format)."""
        # Use same method as XLSX - pandas handles both
        return self._parse_xlsx(filepath)
    
    def _parse_txt(self, filepath: str) -> str:
        """Parse plain text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(filepath, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
                    
            # If all encodings fail, read as binary and decode with errors='ignore'
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
                
        except Exception as e:
            raise IOError(f"Error parsing text file: {str(e)}")
    
    def get_document_info(self, filepath: str) -> Dict[str, Any]:
        """
        Get metadata information about a document.
        
        Args:
            filepath: Path to the document
            
        Returns:
            Dictionary with document metadata
        """
        if not os.path.exists(filepath):
            raise IOError(f"File not found: {filepath}")
            
        info = {
            'filename': os.path.basename(filepath),
            'filepath': filepath,
            'size': os.path.getsize(filepath),
            'extension': os.path.splitext(filepath)[1].lower(),
            'modified': os.path.getmtime(filepath)
        }
        
        return info
    
    def _parse_pdf_with_ocr(self, filepath: str) -> str:
        """Parse PDF using OCR (Optical Character Recognition)."""
        if not OCR_AVAILABLE:
            raise IOError("OCR libraries not available")
        
        try:
            # Set up poppler path for pdf2image
            current_dir = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
            poppler_path = os.path.join(current_dir, 'poppler', 'poppler-23.01.0', 'Library', 'bin')
            
            # Convert PDF to images
            print("Converting PDF pages to images...")
            if os.path.exists(poppler_path):
                print(f"Using local poppler binaries: {poppler_path}")
                images = convert_from_path(filepath, dpi=300, fmt='png', poppler_path=poppler_path)
            else:
                print("Using system poppler (if available)")
                images = convert_from_path(filepath, dpi=300, fmt='png')
            
            text_content = []
            total_chars = 0
            
            for page_num, image in enumerate(images):
                print(f"OCR processing page {page_num + 1}/{len(images)}...")
                
                # Perform OCR on the image
                page_text = pytesseract.image_to_string(image, lang='eng')
                
                if page_text.strip():
                    text_content.append(page_text)
                    chars_on_page = len(page_text.strip())
                    total_chars += chars_on_page
                    print(f"OCR - Page {page_num + 1} extracted {chars_on_page} characters")
                else:
                    print(f"OCR - Page {page_num + 1} extracted 0 characters")
            
            if text_content:
                full_text = "\n\n".join(text_content)
                print(f"OCR successfully extracted {total_chars} characters total")
                
                # Add OCR header to indicate source
                result = f"""[OCR EXTRACTED TEXT]
This text was extracted using Optical Character Recognition (OCR) from scanned images.
OCR accuracy may vary depending on image quality and text clarity.

File: {os.path.basename(filepath)}
Pages processed: {len(images)}
Total characters extracted: {total_chars}

--- EXTRACTED CONTENT ---

{full_text}"""
                return result
            else:
                raise IOError("OCR could not extract any text from the images")
                
        except Exception as e:
            raise IOError(f"OCR processing failed: {str(e)}")
    
    def is_supported(self, filepath: str) -> bool:
        """Check if file format is supported."""
        _, ext = os.path.splitext(filepath.lower())
        return ext in self.supported_formats